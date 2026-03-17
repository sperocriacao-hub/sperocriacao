import os
import uuid
import shutil
from datetime import datetime
from fastapi import APIRouter, File, UploadFile, Form, HTTPException, Depends
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
import fitz
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.database import get_db

router = APIRouter(
    prefix="/api/v1/documents",
    tags=["documents"]
)

UPLOAD_DIR = "storage/uploads"
CONVERTED_DIR = "storage/converted"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(CONVERTED_DIR, exist_ok=True)


@router.post("/convert")
async def convert_document(
    user_id: str = Form(...),
    client: str = Form(...),
    document_type: str = Form(...),
    title: str = Form(...),
    reference: str = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Apenas arquivos PDF são suportados.")

    doc_id = str(uuid.uuid4())
    
    # 1. Salvar o PDF de entrada
    pdf_filename = f"{doc_id}_{file.filename}"
    pdf_path = os.path.join(UPLOAD_DIR, pdf_filename)
    
    with open(pdf_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    # 2. Definir o destino do DOCX
    docx_filename = f"{doc_id}.docx"
    docx_path = os.path.join(CONVERTED_DIR, docx_filename)
    
    start_time = datetime.now()
    
    try:
        # Novo Fluxo: Ler PDF com PyMuPDF, extrair texto puro, gerar DOCX de coluna única.
        doc_pdf = fitz.open(pdf_path)
        
        # Cria um novo documento Word em branco
        doc_word = Document()
        
        # Opcional: ajustar margens para dar mais largura a tabela
        sections = doc_word.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Adiciona a tabela 1 Coluna / 0 Linhas
        table = doc_word.add_table(rows=0, cols=1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for page_num in range(len(doc_pdf)):
            page = doc_pdf.load_page(page_num)
            blocks = page.get_text("blocks")
            
            # Ordenar blocos da esquerda para a direita, cima para baixo
            blocks.sort(key=lambda b: (b[1], b[0]))
            
            for b in blocks:
                text = b[4].strip()
                if not text:
                    continue
                
                # Para evitar multiplos newlines dentro do mesmo bloco prejudicando o visual
                # substitui as quebras de linha dentro do bloco por espaços simples.
                clean_text = " ".join(text.split())
                
                if clean_text:
                    # Adiciona uma nova linha à tabela
                    row_cells = table.add_row().cells
                    cell = row_cells[0]
                    # Insere o parágrafo na célula
                    p = cell.paragraphs[0]
                    p.add_run(clean_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Estilo leve para não ficar colado às margens da célula
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.space_before = Pt(6)

        doc_pdf.close()
        # Salva o resultado no local designado
        doc_word.save(docx_path)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        from app.supabase_client import supabase, upload_to_storage
        
        if supabase:
            try:
                # Faz o upload dos arquivos para o Supabase Storage (Buckets)
                # O bucket name precisa ser criado no Dashboard ("documents")
                await upload_to_storage("documents", pdf_path, f"pdfs/{pdf_filename}", "application/pdf")
                await upload_to_storage("documents", docx_path, f"docx/{docx_filename}", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                # Salvar no Supabase (Base de Dados)
                supabase.table("document_conversions").insert({
                    "id": doc_id,
                    "user_id": user_id,
                    "client": client,
                    "document_type": document_type,
                    "title": title,
                    "reference": reference,
                    "original_pdf_path": pdf_path,
                    "converted_docx_path": docx_path,
                    "status": "COMPLETED",
                    "metadata": {
                        "processing_time_sec": processing_time,
                        "original_filename": file.filename
                    }
                }).execute()
            except Exception as db_err:
                print(f"Aviso: Erro ao salvar no Supabase ou Storage: {db_err}")
                
        # Limpar os ficheiros fisicos do HD efémero para economizar espaco no Render
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        if os.path.exists(docx_path):
            os.remove(docx_path)
        
        return {
            "status": "success",
            "message": "Conversão realizada com sucesso.",
            "data": {
                "document_id": doc_id,
                "preview_url": f"/api/v1/documents/preview/{doc_id}",
                "download_url": f"/api/v1/documents/download/{doc_id}",
                "processing_time": processing_time
            }
        }
        
    except Exception as e:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        if os.path.exists(docx_path):
            os.remove(docx_path)
        raise HTTPException(status_code=500, detail=f"Erro na conversão: {str(e)}")


@router.get("/preview/{doc_id}")
async def get_preview(doc_id: str):
    from app.supabase_client import supabase
    from fastapi.responses import Response

    if supabase:
        try:
            # Baixa o ficheiro binário diretamente do bucket de forma síncrona/segura
            file_data = supabase.storage.from_("documents").download(f"docx/{doc_id}.docx")
            if file_data:
                return Response(
                    content=file_data,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={doc_id}.docx"}
                )
        except Exception as strg_err:
            print(f"Supabase storage miss ou erro: {strg_err}")
            
    # 2. Fallback caso o ficheiro esteja no disco local momentaneamente ou haja erro
    docx_path = os.path.join(CONVERTED_DIR, f"{doc_id}.docx")
    if not os.path.exists(docx_path):
        raise HTTPException(status_code=404, detail="Documento não encontrado na nuvem nem no disco (limite efêmero).")
    
    return FileResponse(
        docx_path, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{doc_id}.docx"
    )

@router.get("/download/{doc_id}")
async def download_converted(doc_id: str):
    return await get_preview(doc_id)
